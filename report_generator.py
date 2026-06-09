import pandas as pd
import numpy as np
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime, timedelta
import io
import os
import base64

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm, inch
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        Image, PageBreak, KeepTogether
    )
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import plotly.graph_objects as go
    import plotly.express as px
    PLOTLY_AVAILABLE = True
except ImportError:
    PLOTLY_AVAILABLE = False

try:
    import kaleido
    KALEIDO_AVAILABLE = True
except ImportError:
    KALEIDO_AVAILABLE = False

from data_processor import (
    aggregate_by_time_granularity, calculate_growth_rate,
    detect_seasonality, detect_anomalies, filter_data, get_available_dimensions
)
from trend_charts import COLOR_PALETTE
from promo_evaluation import (
    evaluate_promo_activity, get_activity_list, DEFAULT_PRODUCT_COST_RATIO
)
from customer_analysis import run_customer_analysis


REPORT_TEMPLATES = {
    'daily': {
        'name': '日报模板',
        'description': '每日销售数据简报，包含核心指标、当日活动表现和异常提醒',
        'sections': ['销售概览', '促销效果', '异常提醒'],
        'default_charts': ['销售趋势', '活动对比']
    },
    'weekly': {
        'name': '周报模板',
        'description': '周度销售分析报告，包含周环比趋势、促销活动汇总和客户洞察',
        'sections': ['销售概览', '趋势分析', '促销效果', '客户洞察'],
        'default_charts': ['销售趋势', '活动对比', '客户分群']
    },
    'monthly': {
        'name': '月报模板',
        'description': '月度经营分析报告，包含完整的销售分析、趋势、促销、客户和改进建议',
        'sections': ['销售概览', '趋势分析', '促销效果', '客户洞察', '改进建议'],
        'default_charts': ['销售趋势', '活动对比', '客户分群', 'RFM雷达图']
    }
}


def _find_chinese_font():
    font_paths = [
        'C:/Windows/Fonts/msyh.ttc',
        'C:/Windows/Fonts/msyh.ttf',
        'C:/Windows/Fonts/simhei.ttf',
        'C:/Windows/Fonts/simsun.ttc',
        '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc',
        '/System/Library/Fonts/PingFang.ttc',
    ]
    for path in font_paths:
        if os.path.exists(path):
            return path
    return None


def _register_chinese_font():
    if not REPORTLAB_AVAILABLE:
        return
    font_path = _find_chinese_font()
    if font_path:
        try:
            pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
            return 'ChineseFont'
        except Exception:
            pass
    return 'Helvetica'


def prepare_report_data(
    df: pd.DataFrame,
    template_type: str = 'weekly',
    time_range: Optional[Dict] = None,
    options: Optional[Dict] = None
) -> Dict[str, Any]:
    options = options or {}
    result = {
        'template_type': template_type,
        'template_info': REPORT_TEMPLATES.get(template_type, REPORT_TEMPLATES['weekly']),
        'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'time_range': time_range or {},
        'options': options,
        'sections': {}
    }

    if df is None or df.empty:
        return result

    work_df = df.copy()
    result['data_rows'] = len(work_df)

    result['sections']['销售概览'] = _build_sales_overview(work_df, options)

    if '趋势分析' in result['template_info']['sections']:
        result['sections']['趋势分析'] = _build_trend_analysis(work_df, options)

    if '促销效果' in result['template_info']['sections']:
        result['sections']['促销效果'] = _build_promo_analysis(work_df, options)

    if '客户洞察' in result['template_info']['sections']:
        result['sections']['客户洞察'] = _build_customer_insight(work_df, options)

    if '异常提醒' in result['template_info']['sections']:
        result['sections']['异常提醒'] = _build_anomaly_alert(work_df, options)

    if '改进建议' in result['template_info']['sections']:
        result['sections']['改进建议'] = _build_improvement_suggestions(work_df, result)

    return result


def _build_sales_overview(df: pd.DataFrame, options: Dict) -> Dict[str, Any]:
    overview = {}

    if '销售额' in df.columns:
        total_sales = float(df['销售额'].sum())
        avg_sales = float(df['销售额'].mean())
        max_sales = float(df['销售额'].max())
        overview['总销售额'] = round(total_sales, 2)
        overview['平均销售额'] = round(avg_sales, 2)
        overview['最高销售额'] = round(max_sales, 2)
        overview['数据周期'] = len(df)

        if '期间' in df.columns and len(df) >= 2:
            sorted_df = df.sort_values('期间').reset_index(drop=True)
            mid = len(sorted_df) // 2
            first_half = sorted_df.iloc[:mid]['销售额'].mean()
            second_half = sorted_df.iloc[mid:]['销售额'].mean()
            if first_half > 0:
                overview['环比变化率(%)'] = round((second_half - first_half) / first_half * 100, 2)
            else:
                overview['环比变化率(%)'] = 0.0

    if '类型' in df.columns:
        type_stats = df.groupby('类型')['销售额'].agg(['sum', 'count']).reset_index()
        type_stats.columns = ['类型', '销售额', '记录数']
        overview['类型分布'] = type_stats.to_dict('records')

    if '产品类别' in df.columns:
        cat_stats = df.groupby('产品类别')['销售额'].sum().sort_values(ascending=False)
        overview['品类销售Top5'] = cat_stats.head(5).reset_index().to_dict('records')

    if '地区' in df.columns:
        reg_stats = df.groupby('地区')['销售额'].sum().sort_values(ascending=False)
        overview['地区销售Top5'] = reg_stats.head(5).reset_index().to_dict('records')

    return overview


def _build_trend_analysis(df: pd.DataFrame, options: Dict) -> Dict[str, Any]:
    analysis = {}
    time_col = '期间' if '期间' in df.columns else df.columns[0]

    for granularity in ['日', '周', '月']:
        try:
            agg_df = aggregate_by_time_granularity(df, time_col, '销售额', granularity)
            if len(agg_df) >= 2:
                growth = calculate_growth_rate(agg_df)
                seasonality = detect_seasonality(agg_df)
                anomalies = detect_anomalies(agg_df)
                analysis[granularity] = {
                    'data': agg_df.to_dict('records'),
                    'growth': growth,
                    'seasonality': seasonality,
                    'anomalies': anomalies
                }
        except Exception:
            continue

    return analysis


def _build_promo_analysis(df: pd.DataFrame, options: Dict) -> Dict[str, Any]:
    result = {}
    activities = get_activity_list(df)
    result['活动数量'] = len(activities)
    result['活动列表'] = activities

    all_results = {}
    for act in activities[:5]:
        try:
            eval_result = evaluate_promo_activity(df, act)
            all_results[act] = {
                '综合评分': eval_result['综合评分']['综合评分'],
                '评级': eval_result['综合评分']['评级'],
                'ROI(%)': eval_result['投资回报分析']['投资回报率(%)'],
                '销售提升率(%)': eval_result['增量销售分析']['销售提升率(%)'],
                '净利率(%)': eval_result['利润边际分析']['净利率(%)'],
                '雷达图数据': eval_result['雷达图数据']
            }
        except Exception:
            continue

    result['活动评估结果'] = all_results

    if all_results:
        sorted_acts = sorted(all_results.items(), key=lambda x: x[1]['综合评分'], reverse=True)
        result['最佳活动'] = sorted_acts[0] if sorted_acts else None
        result['活动排名'] = [
            {'活动名称': k, '综合评分': v['综合评分'], '评级': v['评级']}
            for k, v in sorted_acts
        ]

    return result


def _build_customer_insight(df: pd.DataFrame, options: Dict) -> Dict[str, Any]:
    result = {}
    try:
        n_clusters = options.get('n_clusters', 4)
        cust_result = run_customer_analysis(df, n_clusters=n_clusters)
        if cust_result.get('success'):
            result['总客户数'] = cust_result.get('total_customers', 0)
            result['总交易数'] = cust_result.get('total_transactions', 0)
            result['聚类数量'] = cust_result.get('cluster_info', {}).get('n_clusters', n_clusters)
            result['轮廓系数'] = cust_result.get('cluster_info', {}).get('silhouette_score')
            result['客户群体画像'] = cust_result.get('cluster_profiles', [])
            result['客户分群分布'] = cust_result.get('segment_distribution', [])

            rfm_data = cust_result.get('rfm_data')
            if rfm_data is not None and not rfm_data.empty:
                result['RFM统计'] = {
                    '平均Recency': round(float(rfm_data['Recency'].mean()), 1),
                    '平均Frequency': round(float(rfm_data['Frequency'].mean()), 1),
                    '平均Monetary': round(float(rfm_data['Monetary'].mean()), 2),
                    '平均客单价': round(float(rfm_data['AvgOrderValue'].mean()), 2)
                }
    except Exception as e:
        result['error'] = str(e)

    return result


def _build_anomaly_alert(df: pd.DataFrame, options: Dict) -> Dict[str, Any]:
    result = {'alerts': []}

    if '销售额' in df.columns:
        q1 = df['销售额'].quantile(0.25)
        q3 = df['销售额'].quantile(0.75)
        iqr = q3 - q1
        lower = q1 - 1.5 * iqr
        upper = q3 + 1.5 * iqr

        anomalies_low = df[df['销售额'] < lower]
        anomalies_high = df[df['销售额'] > upper]

        if len(anomalies_low) > 0:
            result['alerts'].append({
                '类型': '销售偏低异常',
                '数量': len(anomalies_low),
                '影响金额': round(float(anomalies_low['销售额'].sum()), 2)
            })
        if len(anomalies_high) > 0:
            result['alerts'].append({
                '类型': '销售偏高异常',
                '数量': len(anomalies_high),
                '影响金额': round(float(anomalies_high['销售额'].sum()), 2)
            })

    if not result['alerts']:
        result['alerts'].append({
            '类型': '正常',
            '数量': 0,
            '影响金额': 0,
            '说明': '未检测到明显的数据异常，整体表现平稳'
        })

    return result


def _build_improvement_suggestions(df: pd.DataFrame, report_data: Dict) -> List[str]:
    suggestions = []

    overview = report_data['sections'].get('销售概览', {})
    growth_rate = overview.get('环比变化率(%)', 0)
    if growth_rate < 0:
        suggestions.append(
            f'销售环比下降 {abs(growth_rate):.1f}%，建议排查原因，考虑推出针对性促销活动刺激消费。'
        )
    elif growth_rate < 5:
        suggestions.append(
            f'销售增长平缓（环比 +{growth_rate:.1f}%），建议优化营销策略，提升转化效率。'
        )
    else:
        suggestions.append(
            f'销售增长良好（环比 +{growth_rate:.1f}%），建议总结经验，持续扩大优势。'
        )

    promo = report_data['sections'].get('促销效果', {})
    rankings = promo.get('活动排名', [])
    if rankings:
        best = rankings[0]
        worst = rankings[-1]
        if len(rankings) >= 2 and worst['综合评分'] < 60:
            suggestions.append(
                f'活动「{worst["活动名称"]}」表现不佳（评分 {worst["综合评分"]:.1f}），'
                f'建议参考最佳活动「{best["活动名称"]}」的策略进行优化。'
            )

    customer = report_data['sections'].get('客户洞察', {})
    profiles = customer.get('客户群体画像', [])
    if profiles:
        low_value = [p for p in profiles if '流失' in p.get('聚类名称', '')]
        if low_value:
            suggestions.append(
                f'检测到 {len(low_value)} 个流失风险客户群体，建议启动客户召回计划和会员忠诚度活动。'
            )

    if not suggestions:
        suggestions.append('当前各项指标表现正常，建议持续监控并优化运营细节。')

    return suggestions


def generate_report_charts(report_data: Dict) -> Dict[str, bytes]:
    charts = {}
    if not PLOTLY_AVAILABLE or not KALEIDO_AVAILABLE:
        return charts

    overview = report_data['sections'].get('销售概览', {})
    trend = report_data['sections'].get('趋势分析', {})
    promo = report_data['sections'].get('促销效果', {})
    customer = report_data['sections'].get('客户洞察', {})

    if trend.get('日'):
        try:
            trend_data = pd.DataFrame(trend['日']['data'])
            fig = go.Figure()
            fig.add_trace(go.Scatter(
                x=trend_data['时间'],
                y=trend_data['销售额'],
                mode='lines+markers',
                name='销售额',
                line=dict(color='#2E86AB', width=3),
                marker=dict(size=8)
            ))
            fig.update_layout(
                title='销售趋势图',
                xaxis_title='时间',
                yaxis_title='销售额',
                plot_bgcolor='white',
                title_x=0.5,
                height=400
            )
            charts['sales_trend'] = fig.to_image(format='png', width=800, height=400, scale=2)
        except Exception:
            pass

    if promo.get('活动排名'):
        try:
            rankings = promo['活动排名']
            fig = go.Figure()
            fig.add_trace(go.Bar(
                x=[r['活动名称'] for r in rankings],
                y=[r['综合评分'] for r in rankings],
                marker_color=['#2E86AB' if r['综合评分'] >= 75 else (
                    '#F18F01' if r['综合评分'] >= 60 else '#e74c3c')
                    for r in rankings],
                text=[f"{r['综合评分']:.1f}" for r in rankings],
                textposition='outside'
            ))
            fig.update_layout(
                title='促销活动综合评分对比',
                xaxis_title='活动名称',
                yaxis_title='综合评分',
                yaxis=dict(range=[0, 100]),
                plot_bgcolor='white',
                title_x=0.5,
                height=400
            )
            charts['promo_comparison'] = fig.to_image(format='png', width=800, height=400, scale=2)
        except Exception:
            pass

    if customer.get('客户分群分布'):
        try:
            segments = customer['客户分群分布']
            fig = go.Figure()
            pie_colors = ['#2E86AB', '#A23B72', '#F18F01', '#C73E1D', '#27ae60', '#8e44ad', '#16a085', '#e67e22']
            fig.add_trace(go.Pie(
                labels=[s['客户分群'] for s in segments],
                values=[s['客户数量'] for s in segments],
                textinfo='label+percent',
                marker=dict(colors=pie_colors[:len(segments)]),
                hole=0.4
            ))
            fig.update_layout(
                title='客户分群分布',
                title_x=0.5,
                height=400
            )
            charts['customer_segment'] = fig.to_image(format='png', width=800, height=400, scale=2)
        except Exception:
            pass

    if customer.get('客户群体画像'):
        try:
            profiles = customer['客户群体画像']
            metrics = ['平均Recency(天)', '平均购买频次', '平均消费总额', '平均客单价']
            metric_labels = ['最近购买', '购买频次', '消费总额', '客单价']
            z = []
            y_labels = []
            for p in profiles:
                y_labels.append(p['聚类名称'])
                row = []
                for m in metrics:
                    val = p.get(m, 0)
                    if isinstance(val, (int, float)):
                        row.append(float(val))
                    else:
                        row.append(0)
                z.append(row)

            fig = go.Figure(data=go.Heatmap(
                z=z,
                x=metric_labels,
                y=y_labels,
                colorscale='Blues',
                text=[[f'{v:.1f}' for v in row] for row in z],
                texttemplate='%{text}'
            ))
            fig.update_layout(
                title='各客户群体特征热力图',
                title_x=0.5,
                height=400
            )
            charts['customer_heatmap'] = fig.to_image(format='png', width=800, height=400, scale=2)
        except Exception:
            pass

    return charts


def generate_pdf_report(report_data: Dict, charts: Optional[Dict[str, bytes]] = None) -> bytes:
    if not REPORTLAB_AVAILABLE:
        raise ImportError('reportlab library is required for PDF generation')

    charts = charts or {}
    buffer = io.BytesIO()
    font_name = _register_chinese_font()

    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        rightMargin=2 * cm, leftMargin=2 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle', parent=styles['Heading1'],
        fontName=font_name, fontSize=22, alignment=1,
        spaceAfter=20, textColor=colors.HexColor('#2E86AB')
    )
    h2_style = ParagraphStyle(
        'CustomH2', parent=styles['Heading2'],
        fontName=font_name, fontSize=16,
        spaceBefore=15, spaceAfter=10,
        textColor=colors.HexColor('#333333')
    )
    h3_style = ParagraphStyle(
        'CustomH3', parent=styles['Heading3'],
        fontName=font_name, fontSize=13,
        spaceBefore=10, spaceAfter=8,
        textColor=colors.HexColor('#2E86AB')
    )
    normal_style = ParagraphStyle(
        'CustomNormal', parent=styles['Normal'],
        fontName=font_name, fontSize=10.5, leading=16,
        spaceAfter=6
    )
    info_style = ParagraphStyle(
        'CustomInfo', parent=styles['Normal'],
        fontName=font_name, fontSize=9,
        textColor=colors.gray, alignment=1
    )

    story = []

    template_info = report_data.get('template_info', {})
    story.append(Paragraph(template_info.get('name', '销售分析报告'), title_style))
    story.append(Paragraph(f'生成时间：{report_data.get("generated_at", "")}', info_style))
    story.append(Spacer(1, 0.5 * cm))

    overview = report_data['sections'].get('销售概览', {})
    if overview:
        story.append(Paragraph('一、销售概览', h2_style))

        kpi_data = []
        kpi_row = []
        for label in ['总销售额', '平均销售额', '最高销售额', '数据周期']:
            if label in overview:
                kpi_row.append(f'{label}: {overview[label]}')
        if kpi_row:
            kpi_data.append(kpi_row)

        if '环比变化率(%)' in overview:
            rate = overview['环比变化率(%)']
            arrow = '↑' if rate >= 0 else '↓'
            color = 'green' if rate >= 0 else 'red'
            kpi_data.append([f'环比变化: {arrow} {abs(rate):.1f}%'])

        if kpi_data:
            kpi_table = Table(kpi_data, colWidths=[4 * cm] * 4)
            kpi_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, -1), font_name),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f8f9fa')),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ]))
            story.append(kpi_table)
            story.append(Spacer(1, 0.3 * cm))

        if overview.get('类型分布'):
            story.append(Paragraph('销售类型分布', h3_style))
            type_rows = [['类型', '销售额', '记录数']]
            for item in overview['类型分布']:
                type_rows.append([
                    str(item.get('类型', '')),
                    f"{item.get('销售额', 0):,.2f}",
                    str(item.get('记录数', 0))
                ])
            type_table = Table(type_rows, colWidths=[5 * cm, 4 * cm, 4 * cm])
            type_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, -1), font_name),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2E86AB')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            story.append(type_table)
            story.append(Spacer(1, 0.3 * cm))

        if charts.get('sales_trend'):
            story.append(Paragraph('销售趋势图', h3_style))
            img_stream = io.BytesIO(charts['sales_trend'])
            story.append(Image(img_stream, width=16 * cm, height=8 * cm))
            story.append(Spacer(1, 0.3 * cm))

    trend = report_data['sections'].get('趋势分析', {})
    if trend:
        story.append(Paragraph('二、趋势分析', h2_style))

        for gran in ['日', '周', '月']:
            gran_data = trend.get(gran)
            if gran_data:
                story.append(Paragraph(f'{gran}度趋势分析', h3_style))
                growth = gran_data.get('growth', {})
                trend_text = f"累计增长率: {growth.get('total_growth_rate', 'N/A')}% | " \
                             f"平均环比: {growth.get('avg_growth_rate', 'N/A')}% | " \
                             f"趋势: {growth.get('trend', 'N/A')}"
                story.append(Paragraph(trend_text, normal_style))

                seasonality = gran_data.get('seasonality', {})
                if seasonality.get('has_seasonality'):
                    story.append(Paragraph(
                        f"季节性模式: {seasonality.get('seasonal_pattern', '检测到季节性特征')}",
                        normal_style
                    ))

                anomalies = gran_data.get('anomalies', [])
                if anomalies:
                    story.append(Paragraph(f"异常检测: 发现 {len(anomalies)} 个异常点", normal_style))

                story.append(Spacer(1, 0.2 * cm))

    promo = report_data['sections'].get('促销效果', {})
    if promo:
        story.append(Paragraph('三、促销效果分析', h2_style))

        story.append(Paragraph(f"活动总数: {promo.get('活动数量', 0)}", normal_style))

        if promo.get('活动排名'):
            story.append(Paragraph('活动综合评分排名', h3_style))
            ranking_rows = [['排名', '活动名称', '综合评分', '评级']]
            for i, r in enumerate(promo['活动排名'], 1):
                ranking_rows.append([
                    str(i), str(r['活动名称']),
                    f"{r['综合评分']:.1f}", str(r['评级'])
                ])
            rank_table = Table(ranking_rows, colWidths=[2 * cm, 7 * cm, 3 * cm, 4 * cm])
            rank_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, -1), font_name),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2E86AB')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            story.append(rank_table)
            story.append(Spacer(1, 0.3 * cm))

        if charts.get('promo_comparison'):
            img_stream = io.BytesIO(charts['promo_comparison'])
            story.append(Image(img_stream, width=16 * cm, height=8 * cm))
            story.append(Spacer(1, 0.3 * cm))

    customer = report_data['sections'].get('客户洞察', {})
    if customer and 'error' not in customer:
        story.append(Paragraph('四、客户洞察分析', h2_style))

        kpi_items = [
            ('总客户数', customer.get('总客户数', 0)),
            ('总交易数', customer.get('总交易数', 0)),
            ('聚类群体数', customer.get('聚类数量', 0)),
        ]
        silhouette = customer.get('轮廓系数')
        if silhouette is not None:
            kpi_items.append(('轮廓系数', f'{silhouette:.4f}'))

        cust_kpi_data = [[f'{k}: {v}' for k, v in kpi_items]]
        cust_kpi_table = Table(cust_kpi_data, colWidths=[4.2 * cm] * len(kpi_items))
        cust_kpi_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), font_name),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f3eafb')),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ]))
        story.append(cust_kpi_table)
        story.append(Spacer(1, 0.3 * cm))

        if customer.get('客户群体画像'):
            story.append(Paragraph('客户群体画像', h3_style))
            profile_rows = [['群体名称', '客户数', '占比(%)', '平均购买频次', '平均客单价']]
            for p in customer['客户群体画像']:
                profile_rows.append([
                    str(p.get('聚类名称', '')),
                    str(p.get('客户数量', 0)),
                    str(p.get('客户占比(%)', 0)),
                    str(p.get('平均购买频次', 0)),
                    f"{p.get('平均客单价', 0):,.0f}"
                ])
            profile_table = Table(profile_rows, colWidths=[4 * cm, 2.5 * cm, 2.5 * cm, 3 * cm, 3.5 * cm])
            profile_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, -1), font_name),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#8e44ad')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('TOPPADDING', (0, 0), (-1, -1), 5),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
            ]))
            story.append(profile_table)
            story.append(Spacer(1, 0.3 * cm))

        if charts.get('customer_segment'):
            img_stream = io.BytesIO(charts['customer_segment'])
            story.append(Image(img_stream, width=16 * cm, height=8 * cm))
            story.append(Spacer(1, 0.3 * cm))

        if charts.get('customer_heatmap'):
            img_stream = io.BytesIO(charts['customer_heatmap'])
            story.append(Image(img_stream, width=16 * cm, height=8 * cm))
            story.append(Spacer(1, 0.3 * cm))

    anomaly = report_data['sections'].get('异常提醒', {})
    if anomaly and anomaly.get('alerts'):
        story.append(Paragraph('五、异常提醒', h2_style))
        for alert in anomaly['alerts']:
            alert_text = f"● {alert.get('类型', '')}: {alert.get('说明', '')}"
            if alert.get('数量', 0) > 0:
                alert_text += f" (共 {alert['数量']} 条, 影响金额: {alert.get('影响金额', 0):,.2f})"
            story.append(Paragraph(alert_text, normal_style))
        story.append(Spacer(1, 0.3 * cm))

    suggestions = report_data['sections'].get('改进建议', [])
    if suggestions:
        story.append(Paragraph('六、改进建议', h2_style))
        for i, s in enumerate(suggestions, 1):
            story.append(Paragraph(f'{i}. {s}', normal_style))
        story.append(Spacer(1, 0.3 * cm))

    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph('—— 报告结束 ——', info_style))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def generate_word_report(report_data: Dict, charts: Optional[Dict[str, bytes]] = None) -> bytes:
    if not DOCX_AVAILABLE:
        raise ImportError('python-docx library is required for Word document generation')

    charts = charts or {}
    buffer = io.BytesIO()
    doc = Document()

    try:
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        style.font.size = Pt(11)
    except Exception:
        pass

    template_info = report_data.get('template_info', {})
    title = doc.add_heading(template_info.get('name', '销售分析报告'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(f'生成时间：{report_data.get("generated_at", "")}')
    info_run.font.size = Pt(9)
    info_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    overview = report_data['sections'].get('销售概览', {})
    if overview:
        doc.add_heading('一、销售概览', level=1)

        kpi_para = doc.add_paragraph()
        kpi_labels = ['总销售额', '平均销售额', '最高销售额', '数据周期']
        kpi_parts = []
        for label in kpi_labels:
            if label in overview:
                kpi_parts.append(f'{label}: {overview[label]}')
        if kpi_parts:
            kpi_run = kpi_para.add_run(' | '.join(kpi_parts))
            kpi_run.bold = True

        if '环比变化率(%)' in overview:
            rate = overview['环比变化率(%)']
            arrow = '↑' if rate >= 0 else '↓'
            trend_para = doc.add_paragraph()
            trend_run = trend_para.add_run(f'环比变化: {arrow} {abs(rate):.1f}%')
            trend_run.font.color.rgb = RGBColor(0x27, 0xae, 0x60) if rate >= 0 else RGBColor(0xe7, 0x4c, 0x3c)
            trend_run.bold = True

        if overview.get('类型分布'):
            doc.add_heading('销售类型分布', level=3)
            table = doc.add_table(rows=1, cols=3)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '类型'
            hdr_cells[1].text = '销售额'
            hdr_cells[2].text = '记录数'
            for item in overview['类型分布']:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item.get('类型', ''))
                row_cells[1].text = f"{item.get('销售额', 0):,.2f}"
                row_cells[2].text = str(item.get('记录数', 0))

        if charts.get('sales_trend'):
            doc.add_heading('销售趋势图', level=3)
            img_stream = io.BytesIO(charts['sales_trend'])
            doc.add_picture(img_stream, width=Inches(6))

    trend = report_data['sections'].get('趋势分析', {})
    if trend:
        doc.add_heading('二、趋势分析', level=1)
        for gran in ['日', '周', '月']:
            gran_data = trend.get(gran)
            if gran_data:
                doc.add_heading(f'{gran}度趋势分析', level=2)
                growth = gran_data.get('growth', {})
                doc.add_paragraph(
                    f"累计增长率: {growth.get('total_growth_rate', 'N/A')}% | "
                    f"平均环比: {growth.get('avg_growth_rate', 'N/A')}% | "
                    f"趋势: {growth.get('trend', 'N/A')}"
                )
                seasonality = gran_data.get('seasonality', {})
                if seasonality.get('has_seasonality'):
                    doc.add_paragraph(f"季节性模式: {seasonality.get('seasonal_pattern', '检测到季节性特征')}")
                anomalies = gran_data.get('anomalies', [])
                if anomalies:
                    doc.add_paragraph(f"异常检测: 发现 {len(anomalies)} 个异常点")

    promo = report_data['sections'].get('促销效果', {})
    if promo:
        doc.add_heading('三、促销效果分析', level=1)
        doc.add_paragraph(f"活动总数: {promo.get('活动数量', 0)}")

        if promo.get('活动排名'):
            doc.add_heading('活动综合评分排名', level=2)
            table = doc.add_table(rows=1, cols=4)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '排名'
            hdr_cells[1].text = '活动名称'
            hdr_cells[2].text = '综合评分'
            hdr_cells[3].text = '评级'
            for i, r in enumerate(promo['活动排名'], 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[1].text = str(r['活动名称'])
                row_cells[2].text = f"{r['综合评分']:.1f}"
                row_cells[3].text = str(r['评级'])

        if charts.get('promo_comparison'):
            img_stream = io.BytesIO(charts['promo_comparison'])
            doc.add_picture(img_stream, width=Inches(6))

    customer = report_data['sections'].get('客户洞察', {})
    if customer and 'error' not in customer:
        doc.add_heading('四、客户洞察分析', level=1)

        cust_kpi_parts = []
        cust_kpi_parts.append(f'总客户数: {customer.get("总客户数", 0)}')
        cust_kpi_parts.append(f'总交易数: {customer.get("总交易数", 0)}')
        cust_kpi_parts.append(f'聚类群体数: {customer.get("聚类数量", 0)}')
        silhouette = customer.get('轮廓系数')
        if silhouette is not None:
            cust_kpi_parts.append(f'轮廓系数: {silhouette:.4f}')
        doc.add_paragraph(' | '.join(cust_kpi_parts))

        if customer.get('客户群体画像'):
            doc.add_heading('客户群体画像', level=2)
            profiles = customer['客户群体画像']
            table = doc.add_table(rows=1, cols=5)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '群体名称'
            hdr_cells[1].text = '客户数'
            hdr_cells[2].text = '占比(%)'
            hdr_cells[3].text = '平均购买频次'
            hdr_cells[4].text = '平均客单价'
            for p in profiles:
                row_cells = table.add_row().cells
                row_cells[0].text = str(p.get('聚类名称', ''))
                row_cells[1].text = str(p.get('客户数量', 0))
                row_cells[2].text = str(p.get('客户占比(%)', 0))
                row_cells[3].text = str(p.get('平均购买频次', 0))
                row_cells[4].text = f"{p.get('平均客单价', 0):,.0f}"

        if charts.get('customer_segment'):
            img_stream = io.BytesIO(charts['customer_segment'])
            doc.add_picture(img_stream, width=Inches(6))

        if charts.get('customer_heatmap'):
            img_stream = io.BytesIO(charts['customer_heatmap'])
            doc.add_picture(img_stream, width=Inches(6))

    anomaly = report_data['sections'].get('异常提醒', {})
    if anomaly and anomaly.get('alerts'):
        doc.add_heading('五、异常提醒', level=1)
        for alert in anomaly['alerts']:
            text = f"● {alert.get('类型', '')}: {alert.get('说明', '')}"
            if alert.get('数量', 0) > 0:
                text += f" (共 {alert['数量']} 条, 影响金额: {alert.get('影响金额', 0):,.2f})"
            doc.add_paragraph(text)

    suggestions = report_data['sections'].get('改进建议', [])
    if suggestions:
        doc.add_heading('六、改进建议', level=1)
        for i, s in enumerate(suggestions, 1):
            doc.add_paragraph(f'{i}. {s}', style='List Number')

    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def get_report_preview_html(report_data: Dict) -> str:
    sections_html = []

    overview = report_data['sections'].get('销售概览', {})
    if overview:
        kpi_html = ''
        for label in ['总销售额', '平均销售额', '最高销售额', '数据周期']:
            if label in overview:
                kpi_html += f'''
                <div style="flex:1; text-align:center; padding:15px; background:#eaf4fb; border-radius:8px; margin:5px;">
                    <div style="color:#666; font-size:13px;">{label}</div>
                    <div style="font-size:24px; font-weight:bold; color:#2E86AB;">{overview[label]}</div>
                </div>'''

        rate_html = ''
        if '环比变化率(%)' in overview:
            rate = overview['环比变化率(%)']
            color = '#27ae60' if rate >= 0 else '#e74c3c'
            arrow = '↑' if rate >= 0 else '↓'
            rate_html = f'''
            <div style="text-align:center; padding:10px; margin-top:10px;">
                <span style="font-weight:bold; color:{color}; font-size:18px;">环比变化: {arrow} {abs(rate):.1f}%</span>
            </div>'''

        table_html = ''
        if overview.get('类型分布'):
            rows = ''.join(
                f'<tr><td>{t.get("类型","")}</td><td>{t.get("销售额",0):,.2f}</td><td>{t.get("记录数",0)}</td></tr>'
                for t in overview['类型分布']
            )
            table_html = f'''
            <h4 style="color:#2E86AB; margin-top:15px;">销售类型分布</h4>
            <table style="width:100%; border-collapse:collapse;">
                <tr style="background:#2E86AB; color:white;">
                    <th style="padding:10px; border:1px solid #ddd;">类型</th>
                    <th style="padding:10px; border:1px solid #ddd;">销售额</th>
                    <th style="padding:10px; border:1px solid #ddd;">记录数</th>
                </tr>
                {rows}
            </table>'''

        sections_html.append(f'''
        <div class="report-section">
            <h2 style="color:#333; border-bottom:2px solid #2E86AB; padding-bottom:8px;">一、销售概览</h2>
            <div style="display:flex; flex-wrap:wrap; gap:10px;">{kpi_html}</div>
            {rate_html}
            {table_html}
        </div>''')

    trend = report_data['sections'].get('趋势分析', {})
    if trend:
        trend_items = ''
        for gran in ['日', '周', '月']:
            gran_data = trend.get(gran)
            if gran_data:
                growth = gran_data.get('growth', {})
                trend_items += f'''
                <div style="background:#f8f9fa; padding:12px; border-radius:8px; margin-bottom:10px;">
                    <h4 style="margin:0 0 8px 0; color:#333;">{gran}度趋势</h4>
                    <div style="color:#555;">
                        累计增长率: {growth.get('total_growth_rate', 'N/A')}% |
                        平均环比: {growth.get('avg_growth_rate', 'N/A')}% |
                        趋势: {growth.get('trend', 'N/A')}
                    </div>
                </div>'''

        sections_html.append(f'''
        <div class="report-section">
            <h2 style="color:#333; border-bottom:2px solid #2E86AB; padding-bottom:8px;">二、趋势分析</h2>
            {trend_items}
        </div>''')

    promo = report_data['sections'].get('促销效果', {})
    if promo:
        ranking_html = ''
        if promo.get('活动排名'):
            rows = ''.join(
                f'<tr><td>{i+1}</td><td>{r["活动名称"]}</td><td>{r["综合评分"]:.1f}</td><td>{r["评级"]}</td></tr>'
                for i, r in enumerate(promo['活动排名'])
            )
            ranking_html = f'''
            <h4 style="color:#2E86AB; margin-top:15px;">活动综合评分排名</h4>
            <table style="width:100%; border-collapse:collapse;">
                <tr style="background:#2E86AB; color:white;">
                    <th style="padding:10px; border:1px solid #ddd;">排名</th>
                    <th style="padding:10px; border:1px solid #ddd;">活动名称</th>
                    <th style="padding:10px; border:1px solid #ddd;">综合评分</th>
                    <th style="padding:10px; border:1px solid #ddd;">评级</th>
                </tr>
                {rows}
            </table>'''

        sections_html.append(f'''
        <div class="report-section">
            <h2 style="color:#333; border-bottom:2px solid #2E86AB; padding-bottom:8px;">三、促销效果分析</h2>
            <p style="color:#555;">活动总数: <strong>{promo.get("活动数量", 0)}</strong></p>
            {ranking_html}
        </div>''')

    customer = report_data['sections'].get('客户洞察', {})
    if customer and 'error' not in customer:
        profile_html = ''
        if customer.get('客户群体画像'):
            rows = ''.join(
                f'<tr><td>{p.get("聚类名称","")}</td><td>{p.get("客户数量",0)}</td>'
                f'<td>{p.get("客户占比(%)",0)}</td><td>{p.get("平均购买频次",0)}</td>'
                f'<td>{p.get("平均客单价",0):,.0f}</td></tr>'
                for p in customer['客户群体画像']
            )
            profile_html = f'''
            <h4 style="color:#8e44ad; margin-top:15px;">客户群体画像</h4>
            <table style="width:100%; border-collapse:collapse;">
                <tr style="background:#8e44ad; color:white;">
                    <th style="padding:10px; border:1px solid #ddd;">群体名称</th>
                    <th style="padding:10px; border:1px solid #ddd;">客户数</th>
                    <th style="padding:10px; border:1px solid #ddd;">占比(%)</th>
                    <th style="padding:10px; border:1px solid #ddd;">平均购买频次</th>
                    <th style="padding:10px; border:1px solid #ddd;">平均客单价</th>
                </tr>
                {rows}
            </table>'''

        sections_html.append(f'''
        <div class="report-section">
            <h2 style="color:#333; border-bottom:2px solid #2E86AB; padding-bottom:8px;">四、客户洞察分析</h2>
            <p style="color:#555;">
                总客户数: <strong>{customer.get("总客户数", 0)}</strong> |
                总交易数: <strong>{customer.get("总交易数", 0)}</strong> |
                聚类群体数: <strong>{customer.get("聚类数量", 0)}</strong>
            </p>
            {profile_html}
        </div>''')

    anomaly = report_data['sections'].get('异常提醒', {})
    if anomaly and anomaly.get('alerts'):
        alert_items = ''.join(
            f'<li style="padding:8px 0; color:#555;">'
            f'<strong>{a.get("类型","")}</strong>: {a.get("说明", "")}'
            + (f' (共 {a["数量"]} 条, 影响金额: {a.get("影响金额", 0):,.2f})' if a.get('数量', 0) > 0 else '')
            + '</li>'
            for a in anomaly['alerts']
        )
        sections_html.append(f'''
        <div class="report-section">
            <h2 style="color:#333; border-bottom:2px solid #2E86AB; padding-bottom:8px;">五、异常提醒</h2>
            <ul style="list-style:none; padding:0;">{alert_items}</ul>
        </div>''')

    suggestions = report_data['sections'].get('改进建议', [])
    if suggestions:
        sug_items = ''.join(
            f'<li style="padding:8px 0; color:#555; line-height:1.6;">{s}</li>'
            for s in suggestions
        )
        sections_html.append(f'''
        <div class="report-section">
            <h2 style="color:#333; border-bottom:2px solid #2E86AB; padding-bottom:8px;">六、改进建议</h2>
            <ol style="padding-left:20px;">{sug_items}</ol>
        </div>''')

    template_info = report_data.get('template_info', {})
    html = f'''
    <div style="font-family: Arial, 'Microsoft YaHei', sans-serif; padding:20px; max-width:900px; margin:0 auto; background:white; border-radius:10px;">
        <h1 style="text-align:center; color:#2E86AB; margin-bottom:5px;">{template_info.get("name", "销售分析报告")}</h1>
        <p style="text-align:center; color:#999; font-size:13px; margin-bottom:25px;">生成时间：{report_data.get("generated_at", "")}</p>
        {''.join(sections_html)}
        <div style="text-align:center; color:#999; margin-top:30px; padding:15px; border-top:1px solid #eee;">—— 报告预览 ——</div>
    </div>'''
    return html
